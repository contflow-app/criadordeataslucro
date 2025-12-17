# app.py — Gerador de ATAs (Individual + Lote) com OCR e extração (Regex + GPT opcional)
# Template: templates/TEMPLATE_ATA.docx (formatado)

from __future__ import annotations

import io
import os
import re
import time
import zipfile
import tempfile
import subprocess
from dataclasses import dataclass, asdict
from datetime import datetime, date
from pathlib import Path
from typing import List, Tuple, Optional, Dict, Any

import streamlit as st
import pdfplumber
from openpyxl import Workbook
from openpyxl.utils import get_column_letter
from docx import Document

# OCR (PDF escaneado)
try:
    import pytesseract
    from pdf2image import convert_from_bytes
    OCR_OK = True
except Exception:
    OCR_OK = False

# OpenAI (opcional)
try:
    from openai import OpenAI
    OPENAI_OK = True
except Exception:
    OPENAI_OK = False


# =========================
# Constantes / padrões
# =========================
DEFAULT_PRESIDENTE = "STANLEY DE SOUZA MOREIRA"
DEFAULT_ADVOGADO = "MARCO AURÉLIO POFFO"

PT_MONTHS = {
    1: "janeiro", 2: "fevereiro", 3: "março", 4: "abril", 5: "maio", 6: "junho",
    7: "julho", 8: "agosto", 9: "setembro", 10: "outubro", 11: "novembro", 12: "dezembro",
}

CNPJ_RE = re.compile(r"\b\d{2}\.?\d{3}\.?\d{3}/?\d{4}-?\d{2}\b")
CPF_RE  = re.compile(r"\b\d{3}\.?\d{3}\.?\d{3}-?\d{2}\b")
NIRE_RE = re.compile(r"\bNIRE\s*(?:n[º°o]?\.?)?\s*[:\-]?\s*([0-9.\-]{5,})", re.IGNORECASE)

# Nome empresarial (heurístico)
RAZAO_RE = re.compile(
    r"([A-ZÁÂÃÉÊÍÓÔÕÚÇ0-9][A-ZÁÂÃÉÊÍÓÔÕÚÇ0-9\s\.\-&]{6,}(?:LTDA|Ltda|S\/A|SA|EIRELI|ME|EPP))"
)

# Endereço: frases típicas
ENDERECO_PATTERNS = [
    re.compile(r"localizad[ao]\s+no\s+endere[cç]o\s+(.+?)(?:\.\s|\n)", re.IGNORECASE | re.DOTALL),
    re.compile(r"tem\s+sede\s*:\s*(.+?)(?:\.\s|\n)", re.IGNORECASE | re.DOTALL),
    re.compile(r"sede\s+na\s+(.+?)(?:\.\s|\n)", re.IGNORECASE | re.DOTALL),
]


# =========================
# Modelos de dados
# =========================
@dataclass
class Socio:
    nome: str
    cpf: str

@dataclass
class Extracao:
    razao_social: str = ""
    cnpj: str = ""
    nire: str = ""
    endereco: str = ""
    cidade_uf: str = ""
    socios: List[Socio] = None

    def __post_init__(self):
        if self.socios is None:
            self.socios = []

@dataclass
class ResultadoProcessamento:
    arquivo_origem: str
    status: str
    mensagem: str
    metodo: str
    ocr_usado: bool
    razao_social: str
    cnpj: str
    nire: str
    cidade_uf: str
    qtd_socios: int
    tempo_ms: int
    docx_bytes: Optional[bytes] = None
    pdf_bytes: Optional[bytes] = None


# =========================
# Utilitários
# =========================
def normalize_spaces(s: str) -> str:
    s = (s or "").replace("\u00A0", " ")
    s = re.sub(r"[ \t]+", " ", s)
    return s.strip()

def safe_first_match(pattern: re.Pattern, text: str, group: int = 0) -> str:
    m = pattern.search(text or "")
    return m.group(group) if m else ""

def extract_text_from_pdf(pdf_bytes: bytes, ocr_if_needed: bool = True, min_chars: int = 250) -> Tuple[str, bool]:
    """
    Retorna (texto, ocr_usado). Primeiro tenta texto nativo. Se insuficiente, faz OCR.
    """
    native = ""
    try:
        parts = []
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            for page in pdf.pages:
                t = (page.extract_text() or "").strip()
                if t:
                    parts.append(t)
        native = "\n".join(parts).strip()
    except Exception:
        native = ""

    if len(native) >= min_chars:
        return native, False

    if not ocr_if_needed or not OCR_OK:
        return native, False

    try:
        images = convert_from_bytes(pdf_bytes, dpi=250)
        ocr_parts = []
        for img in images:
            txt = pytesseract.image_to_string(img, lang="por", config="--oem 3 --psm 6")
            txt = (txt or "").strip()
            if txt:
                ocr_parts.append(txt)
        ocr_text = "\n".join(ocr_parts).strip()
        if len(ocr_text) > len(native):
            return ocr_text, True
        return native, False
    except Exception:
        return native, False

def extract_text_from_docx(docx_bytes: bytes) -> str:
    doc = Document(io.BytesIO(docx_bytes))
    parts = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text)
    return "\n".join(parts)

def city_uf_from_text(text: str) -> str:
    """
    Conservador: só retorna se aparecer explicitamente 'Cidade/UF' ou 'Cidade - UF'
    """
    t = text or ""
    m = re.search(r"\b([A-ZÁÂÃÉÊÍÓÔÕÚÇ][A-Za-zÁÂÃÉÊÍÓÔÕÚÇ\s]{2,})\s*/\s*([A-Z]{2})\b", t)
    if m:
        return f"{normalize_spaces(m.group(1))}/{m.group(2)}"
    m = re.search(r"\b([A-ZÁÂÃÉÊÍÓÔÕÚÇ][A-Za-zÁÂÃÉÊÍÓÔÕÚÇ\s]{2,})\s*-\s*([A-Z]{2})\b", t)
    if m:
        return f"{normalize_spaces(m.group(1))}/{m.group(2)}"
    return ""

def extract_endereco(text: str) -> Tuple[str, str]:
    """
    Retorna (endereco, cidade_uf).
    """
    endereco = ""
    for pat in ENDERECO_PATTERNS:
        m = pat.search(text or "")
        if m:
            endereco = normalize_spaces(m.group(1))
            break
    cidade_uf = city_uf_from_text(endereco) if endereco else ""
    return endereco, cidade_uf

def extract_socios_heuristico(text: str) -> List[Socio]:
    """
    Estratégias:
    1) Padrão assinatura: linha nome, próxima linha CPF
    2) Inline: 'NOME ... CPF nº ...'
    3) Fallback: só CPFs
    Dedup por CPF.
    """
    t = text or ""
    lines = [normalize_spaces(l) for l in t.splitlines() if normalize_spaces(l)]
    socios: List[Socio] = []
    seen = set()

    # (1) assinatura
    for i in range(1, len(lines)):
        if re.search(r"\bCPF\b", lines[i], flags=re.IGNORECASE):
            mcpf = CPF_RE.search(lines[i])
            if not mcpf:
                continue
            cpf = mcpf.group(0)
            nome = lines[i-1]
            if cpf in seen:
                continue
            if nome and len(nome) >= 5 and not re.search(r"\d{2}\.\d{3}\.\d{3}/", nome):
                socios.append(Socio(nome=nome, cpf=cpf))
                seen.add(cpf)

    # (2) inline
    if not socios:
        for m in re.finditer(r"([A-ZÁÂÃÉÊÍÓÔÕÚÇ][A-ZÁÂÃÉÊÍÓÔÕÚÇ\s]{4,120}).{0,120}?\bCPF\b.{0,40}?("+CPF_RE.pattern+r")", t):
            nome = normalize_spaces(m.group(1))
            cpf = m.group(2)
            if cpf not in seen:
                socios.append(Socio(nome=nome, cpf=cpf))
                seen.add(cpf)

    # (3) fallback CPFs
    if not socios:
        for m in CPF_RE.finditer(t):
            cpf = m.group(0)
            if cpf not in seen:
                socios.append(Socio(nome="", cpf=cpf))
                seen.add(cpf)

    return socios

def regex_extract(text: str) -> Extracao:
    t = text or ""
    razao = ""
    m = RAZAO_RE.search(t)
    if m:
        razao = normalize_spaces(m.group(1))[:180]
    cnpj = safe_first_match(CNPJ_RE, t, 0)
    nire = safe_first_match(NIRE_RE, t, 1)
    endereco, cidade_uf = extract_endereco(t)
    socios = extract_socios_heuristico(t)
    return Extracao(
        razao_social=razao,
        cnpj=cnpj,
        nire=nire,
        endereco=endereco,
        cidade_uf=cidade_uf,
        socios=socios
    )

def gpt_extract(text: str, api_key: str) -> Tuple[Optional[Extracao], str]:
    """
    Extração via OpenAI com Structured Outputs (JSON Schema).
    Retorna (Extracao|None, raw_output_text).
    """
    if not OPENAI_OK:
        return None, "Biblioteca openai não instalada."
    if not api_key:
        return None, "OPENAI_API_KEY ausente (secrets)."

    client = OpenAI(api_key=api_key)

    schema = {
      "name": "extract_societario",
      "schema": {
        "type": "object",
        "additionalProperties": False,
        "properties": {
          "razao_social": {"type":"string"},
          "cnpj": {"type":"string"},
          "nire": {"type":"string"},
          "endereco": {"type":"string"},
          "cidade_uf": {"type":"string"},
          "socios": {
            "type":"array",
            "items":{
              "type":"object",
              "additionalProperties": False,
              "properties":{
                "nome":{"type":"string"},
                "cpf":{"type":"string"}
              },
              "required":["nome","cpf"]
            }
          }
        },
        "required":["razao_social","cnpj","nire","endereco","cidade_uf","socios"]
      },
      "strict": True
    }

    prompt = f"""
Extraia dados societários do texto abaixo e retorne APENAS JSON no schema.
Regras:
- Não invente dados.
- Liste TODOS os sócios com NOME e CPF quando disponíveis. Se CPF não aparecer, cpf="".
- Não confunda NIRE com CPF.
- cidade_uf apenas se explícito.
- Se algo não estiver no texto, retorne "".

TEXTO:
\"\"\"{text[:120000]}\"\"\"
"""

    resp = client.responses.create(
        model="gpt-4o-mini",
        input=[{"role":"user","content":prompt}],
        text={"format":{"type":"json_schema","json_schema":schema}}
    )
    raw = getattr(resp, "output_text", "") or ""

    import json
    try:
        data = json.loads(raw)
        socios = [Socio(nome=s.get("nome",""), cpf=s.get("cpf","")) for s in (data.get("socios") or [])]
        ext = Extracao(
            razao_social=data.get("razao_social",""),
            cnpj=data.get("cnpj",""),
            nire=data.get("nire",""),
            endereco=data.get("endereco",""),
            cidade_uf=data.get("cidade_uf",""),
            socios=socios
        )
        return ext, raw
    except Exception as e:
        return None, f"Falha parse JSON: {e}. Raw: {raw[:2000]}"

def docx_to_pdf_via_libreoffice(docx_bytes: bytes) -> Optional[bytes]:
    """
    Conversão DOCX->PDF via LibreOffice (soffice).
    """
    with tempfile.TemporaryDirectory() as tmp:
        docx_path = os.path.join(tmp, "ata.docx")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)
        cmd = ["soffice", "--headless", "--convert-to", "pdf", "--outdir", tmp, docx_path]
        try:
            subprocess.run(cmd, check=True, capture_output=True)
            pdf_path = os.path.join(tmp, "ata.pdf")
            if not os.path.exists(pdf_path):
                return None
            with open(pdf_path, "rb") as f:
                return f.read()
        except Exception:
            return None

def _remove_paragraph(paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)

def fill_template_docx(template_path: str, ext: Extracao, ata_date: date) -> bytes:
    """
    Preenche o template DOCX (já formatado).
    Estratégia: substituir trechos “âncora” e recriar blocos (presença e assinaturas).
    """
    doc = Document(template_path)

    dia = str(ata_date.day)
    mes = PT_MONTHS[ata_date.month]
    ano = str(ata_date.year)

    presidente = DEFAULT_PRESIDENTE
    secretario = DEFAULT_ADVOGADO

    # Cidade/UF padrão = extraída da sede (se houver). Se não, mantém em branco.
    cidade_uf = ext.cidade_uf

    # 1) Substituições em parágrafos-chave (âncoras do template)
    for p in doc.paragraphs:
        t = p.text or ""

        # Cabeçalho com razão social / CNPJ / NIRE / endereço
        if "Sociedade empresária limitada" in t and "CNPJ/MF" in t:
            # empresa
            if ext.razao_social:
                t = re.sub(r"–\s*.+?,\s*inscrita", f"– {ext.razao_social}, inscrita", t)
            # cnpj
            if ext.cnpj:
                t = re.sub(r"CNPJ/MF\s*n[º°]\.\s*[^–]+–", f"CNPJ/MF nº. {ext.cnpj} –", t)
            # nire (se template contiver)
            if "NIRE" in t and ext.nire:
                t = re.sub(r"NIRE:\s*[^ \n]+", f"NIRE: {ext.nire}", t)
            # endereço
            if ext.endereco:
                t = re.sub(r"localizada no endereço\s+_+\.?", f"localizada no endereço {ext.endereco}.", t)
                t = re.sub(r"localizada no endereço\s+.+?\.", f"localizada no endereço {ext.endereco}.", t)
            p.text = t

        # Linha de abertura da reunião (data + sede)
        if t.strip().startswith("Aos dias") and "na sede da" in t:
            t = re.sub(r"Aos\s+dias\s+.+?\s+do\s+mês\s+de\s+.+?\s+do\s+ano\s+de\s+.+?,",
                       f"Aos dias {dia} do mês de {mes} do ano de {ano},", t)
            if ext.razao_social:
                t = re.sub(r"na sede da\s+.+?,\s+localizada", f"na sede da {ext.razao_social}, localizada", t)
            if ext.endereco:
                t = re.sub(r"localizada no endereço\s*\.?", f"localizada no endereço {ext.endereco}.", t)
                t = re.sub(r"localizada no endereço\s+.+?\.", f"localizada no endereço {ext.endereco}.", t)
            p.text = t

        # Mesa (presidente/secretário)
        if t.strip().startswith("DA COMPOSIÇÃO DA MESA"):
            t = re.sub(r"presidida por\s+_+", f"presidida por {presidente}", t)
            t = re.sub(r"SECRETÁRIO\s+_+", f"SECRETÁRIO {secretario}", t)
            p.text = t

        # Linha final (Cidade, ... de ... de AAAA)
        # tenta preencher cidade e ano/mês se houver placeholders
        if re.match(r"^_+,", t.strip()):
            if cidade_uf:
                cidade = cidade_uf.split("/")[0]
                t = re.sub(r"^_+", cidade, t)
            # mês/ano
            t = re.sub(r"de\s+_+\s+de\s+_+\.?$", f"de {mes} de {ano}.", t)
            p.text = t

    # 2) Presença (lista numerada) — substitui tudo entre "DA PRESENÇA" e "Os quais"
    pres_idx = None
    osquais_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("DA PRESENÇA"):
            pres_idx = i
        if pres_idx is not None and p.text.strip().startswith("Os quais"):
            osquais_idx = i
            break

    if pres_idx is not None and osquais_idx is not None:
        # remove linhas numeradas existentes
        to_remove = []
        for j in range(pres_idx + 1, osquais_idx):
            if re.match(r"^\d+\.\s*", doc.paragraphs[j].text.strip()):
                to_remove.append(doc.paragraphs[j])
        for par in to_remove:
            _remove_paragraph(par)

        # inserir após parágrafo "DA PRESENÇA ..."
        from docx.oxml import OxmlElement
        from docx.text.paragraph import Paragraph

        def insert_paragraph_after(paragraph, text):
            new_p = OxmlElement("w:p")
            paragraph._p.addnext(new_p)
            new_para = Paragraph(new_p, paragraph._parent)
            if text:
                new_para.add_run(text)
            return new_para

        cursor = doc.paragraphs[pres_idx]
        socios = ext.socios or []
        if socios:
            for idx, s in enumerate(socios, start=1):
                nome = (s.nome or "").strip() or "________________________"
                cursor = insert_paragraph_after(cursor, f"{idx}. {nome}")
        else:
            for idx in range(1, 4):
                cursor = insert_paragraph_after(cursor, f"{idx}. ________________________")

    # 3) Assinaturas — remove tudo após "ASSINATURAS DOS SÓCIOS:" e recria
    sig_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("ASSINATURAS DOS SÓCIOS"):
            sig_idx = i
            break

    if sig_idx is not None:
        for p in list(doc.paragraphs)[sig_idx + 1:]:
            _remove_paragraph(p)

        from docx.oxml import OxmlElement
        from docx.text.paragraph import Paragraph

        def insert_after(paragraph, text):
            new_p = OxmlElement("w:p")
            paragraph._p.addnext(new_p)
            new_para = Paragraph(new_p, paragraph._parent)
            if text:
                new_para.add_run(text)
            return new_para

        last = doc.paragraphs[sig_idx]
        for s in (ext.socios or []):
            last = insert_after(last, "")
            last = insert_after(last, "_______________________________________")
            nome = (s.nome or "").strip() or "________________________"
            cpf = (s.cpf or "").strip()
            last = insert_after(last, nome)
            if cpf:
                last = insert_after(last, f"CPF nº {cpf}")

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

def make_report_xlsx(rows: List[ResultadoProcessamento]) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "Relatório"
    headers = [
        "arquivo_origem","status","mensagem","ocr_usado","metodo",
        "razao_social","cnpj","nire","cidade_uf","qtd_socios",
        "docx_gerado","pdf_gerado","tempo_ms"
    ]
    ws.append(headers)
    for r in rows:
        ws.append([
            r.arquivo_origem, r.status, r.mensagem,
            "sim" if r.ocr_usado else "não",
            r.metodo,
            r.razao_social, r.cnpj, r.nire, r.cidade_uf, r.qtd_socios,
            "sim" if r.docx_bytes else "não",
            "sim" if r.pdf_bytes else "não",
            r.tempo_ms
        ])
    for col in range(1, len(headers) + 1):
        ws.column_dimensions[get_column_letter(col)].width = 22
    bio = io.BytesIO()
    wb.save(bio)
    bio.seek(0)
    return bio.getvalue()

def load_batch_files(uploaded_files, uploaded_zip) -> List[Tuple[str, bytes]]:
    files: List[Tuple[str, bytes]] = []
    if uploaded_files:
        for f in uploaded_files:
            files.append((f.name, f.read()))
    if uploaded_zip:
        zbytes = uploaded_zip.read()
        with zipfile.ZipFile(io.BytesIO(zbytes)) as z:
            for name in z.namelist():
                if name.endswith("/"):
                    continue
                if name.lower().endswith((".pdf", ".docx")):
                    files.append((Path(name).name, z.read(name)))
    return files

def process_one(name: str, bts: bytes, template_path: str, ata_date: date, try_pdf: bool, use_gpt: bool) -> ResultadoProcessamento:
    t0 = time.time()
    ocr_used = False
    metodo = "regex"
    try:
        # 1) extrai texto
        if name.lower().endswith(".pdf"):
            text, ocr_used = extract_text_from_pdf(bts, ocr_if_needed=True)
        elif name.lower().endswith(".docx"):
            text = extract_text_from_docx(bts)
        else:
            return ResultadoProcessamento(
                arquivo_origem=name, status="ERRO", mensagem="Formato não suportado (use PDF/DOCX).",
                metodo="n/a", ocr_usado=False, razao_social="", cnpj="", nire="", cidade_uf="",
                qtd_socios=0, tempo_ms=int((time.time()-t0)*1000)
            )

        if not text.strip():
            return ResultadoProcessamento(
                arquivo_origem=name, status="ERRO", mensagem="Texto vazio (PDF pode ser imagem e OCR falhou/indisponível).",
                metodo="n/a", ocr_usado=ocr_used, razao_social="", cnpj="", nire="", cidade_uf="",
                qtd_socios=0, tempo_ms=int((time.time()-t0)*1000)
            )

        # 2) extração
        ext: Optional[Extracao] = None
        gpt_raw_msg = ""
        if use_gpt:
            api_key = st.secrets.get("OPENAI_API_KEY", "")
            ext, gpt_raw_msg = gpt_extract(text, api_key)
            if ext:
                metodo = "gpt"
            else:
                metodo = "regex (fallback)"
                ext = regex_extract(text)
        else:
            ext = regex_extract(text)

        # 3) qualidade mínima
        issues = []
        if not ext.razao_social:
            issues.append("Razão social ausente")
        if not ext.cnpj:
            issues.append("CNPJ ausente")
        if not (ext.socios and any((s.nome or "").strip() for s in ext.socios)):
            issues.append("Nomes de sócios ausentes")
        status = "OK" if not issues else "PENDENTE"
        msg = "Gerado" if status == "OK" else "; ".join(issues)
        if ocr_used:
            msg += " (OCR aplicado)"
        if use_gpt and metodo != "gpt":
            msg += " (GPT falhou, fallback regex)"
        if use_gpt and metodo == "gpt":
            msg += " (GPT)"

        # 4) gera DOCX/PDF
        docx_bytes = fill_template_docx(template_path, ext, ata_date)
        pdf_bytes = docx_to_pdf_via_libreoffice(docx_bytes) if try_pdf else None

        tempo_ms = int((time.time() - t0) * 1000)
        return ResultadoProcessamento(
            arquivo_origem=name,
            status=status,
            mensagem=msg,
            metodo=metodo,
            ocr_usado=ocr_used,
            razao_social=ext.razao_social,
            cnpj=ext.cnpj,
            nire=ext.nire,
            cidade_uf=ext.cidade_uf,
            qtd_socios=len(ext.socios or []),
            tempo_ms=tempo_ms,
            docx_bytes=docx_bytes,
            pdf_bytes=pdf_bytes,
        )
    except Exception as e:
        tempo_ms = int((time.time() - t0) * 1000)
        return ResultadoProcessamento(
            arquivo_origem=name, status="ERRO", mensagem=f"{type(e).__name__}: {e}",
            metodo=metodo, ocr_usado=ocr_used, razao_social="", cnpj="", nire="", cidade_uf="",
            qtd_socios=0, tempo_ms=tempo_ms
        )


# =========================
# UI
# =========================
st.set_page_config(page_title="Criador de Atas (Refatorado)", layout="wide")
st.title("Criador de Atas — Refatorado (Template + OCR + GPT opcional)")
st.caption("Usa o template DOCX formatado. Extrai dados do contrato (PDF/DOCX) com OCR automático e opção de GPT para melhorar a extração.")

with st.sidebar:
    st.subheader("Template")
    template_path = st.text_input("Caminho do template", value="templates/TEMPLATE_ATA.docx")
    try_pdf = st.checkbox("Gerar PDF (requer LibreOffice)", value=False)

    st.subheader("Padrões fixos")
    st.text_input("Presidente (fixo)", value=DEFAULT_PRESIDENTE, disabled=True)
    st.text_input("Advogado (fixo)", value=DEFAULT_ADVOGADO, disabled=True)

    st.subheader("Data da ATA")
    ata_date = st.date_input("Data (padrão = hoje)", value=datetime.now().date())

    st.subheader("Extração")
    use_gpt = st.checkbox("Usar GPT para melhorar extração", value=False)
    if use_gpt:
        if not OPENAI_OK:
            st.error("Biblioteca openai não instalada. Verifique requirements.txt.")
        elif not st.secrets.get("OPENAI_API_KEY", ""):
            st.warning("Falta OPENAI_API_KEY em secrets. (Streamlit Cloud → Settings → Secrets)")

tab1, tab2 = st.tabs(["Individual", "Lote"])

with tab1:
    st.subheader("Individual")
    up = st.file_uploader("Suba 1 contrato (PDF ou DOCX)", type=["pdf", "docx"])
    if up:
        bts = up.read()
        r = process_one(up.name, bts, template_path, ata_date, try_pdf, use_gpt)

        st.write(f"**Status:** {r.status}  |  **Método:** {r.metodo}")
        st.write(f"**Mensagem:** {r.mensagem}")
        st.write(f"**Sócios encontrados:** {r.qtd_socios}")

        base = Path(up.name).stem
        if r.docx_bytes:
            st.download_button(
                "Baixar ATA (DOCX)",
                data=r.docx_bytes,
                file_name=f"ATA_{base}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        if r.pdf_bytes:
            st.download_button(
                "Baixar ATA (PDF)",
                data=r.pdf_bytes,
                file_name=f"ATA_{base}.pdf",
                mime="application/pdf"
            )

with tab2:
    st.subheader("Lote")
    uploaded_files = st.file_uploader("Vários arquivos (PDF/DOCX)", type=["pdf", "docx"], accept_multiple_files=True)
    uploaded_zip = st.file_uploader("Ou ZIP com PDFs/DOCXs", type=["zip"])
    batch = load_batch_files(uploaded_files, uploaded_zip)
    st.write(f"Arquivos carregados: **{len(batch)}**")

    continue_on_error = st.checkbox("Continuar mesmo se houver erro", value=True)
    include_excel_inside_zip = st.checkbox("Incluir Excel dentro do ZIP", value=True)

    if st.button("Processar lote", disabled=(len(batch) == 0)):
        rows: List[ResultadoProcessamento] = []
        prog = st.progress(0.0)
        box = st.empty()

        for i, (name, bts) in enumerate(batch, start=1):
            box.write(f"Processando {i}/{len(batch)}: **{name}**")
            r = process_one(name, bts, template_path, ata_date, try_pdf, use_gpt)
            rows.append(r)
            prog.progress(i / len(batch))
            if (not continue_on_error) and r.status == "ERRO":
                st.error(f"Parado no erro: {name} — {r.mensagem}")
                break

        excel_bytes = make_report_xlsx(rows)

        zip_out = io.BytesIO()
        with zipfile.ZipFile(zip_out, "w", zipfile.ZIP_DEFLATED) as z:
            for r in rows:
                base = Path(r.arquivo_origem).stem
                if r.docx_bytes:
                    z.writestr(f"ATAs/{base}.docx", r.docx_bytes)
                if r.pdf_bytes:
                    z.writestr(f"ATAs/{base}.pdf", r.pdf_bytes)
            if include_excel_inside_zip:
                z.writestr("relatorio_processamento.xlsx", excel_bytes)

        zip_out.seek(0)

        ok = sum(1 for r in rows if r.status == "OK")
        pend = sum(1 for r in rows if r.status == "PENDENTE")
        err = sum(1 for r in rows if r.status == "ERRO")
        st.success(f"Concluído. OK: {ok} | PENDENTE: {pend} | ERRO: {err}")

        st.download_button(
            "Baixar Relatório (Excel)",
            data=excel_bytes,
            file_name="relatorio_processamento.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
        st.download_button(
            "Baixar ZIP (ATAs + Excel opcional)",
            data=zip_out.getvalue(),
            file_name="atas_geradas.zip",
            mime="application/zip",
        )

        st.dataframe([{
            "arquivo": r.arquivo_origem,
            "status": r.status,
            "metodo": r.metodo,
            "ocr": "sim" if r.ocr_usado else "não",
            "socios": r.qtd_socios,
            "mensagem": r.mensagem
        } for r in rows])
